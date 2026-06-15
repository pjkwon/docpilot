"""
EBS_영어학습_정리노트_템플릿.docx 를 복사해
{{플레이스홀더}} 문법을 가진 파이프라인용 버전을 생성합니다.
출력: data/EBS_영어학습_정리노트_v2.docx
"""

import copy
from pathlib import Path
from docx import Document

SRC = Path("data/EBS_영어학습_정리노트_템플릿.docx")
DST = Path("data/EBS_영어학습_정리노트_v2.docx")


def set_cell_text(cell, placeholder: str) -> None:
    """셀의 첫 paragraph 첫 run에 텍스트를 설정하고 나머지 run을 비운다."""
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = placeholder
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(placeholder)
    # 추가 paragraph는 비운다
    for extra_para in cell.paragraphs[1:]:
        for run in extra_para.runs:
            run.text = ""


doc = Document(str(SRC))

# ── 본문 단락 수정 ──────────────────────────────────────────
for para in doc.paragraphs:
    text = para.text
    runs = para.runs

    if "YYYY-MM-DD" in text:
        # "학습일: {{학습일}}   |   주제: {{주제}}"
        for run in runs:
            if "YYYY-MM-DD" in run.text:
                run.text = "{{학습일}}"
            elif "오늘의 학습 주제를 입력하세요" in run.text:
                run.text = "{{주제}}"

    elif "대화 상황을 간략히 묘사하세요" in text:
        for run in runs:
            if "대화 상황을 간략히 묘사하세요" in run.text:
                run.text = "{{상황설명}}"

    elif "팁 1:" in text:
        for run in runs:
            if "문화적 맥락 또는 학습 팁을 입력하세요" in run.text:
                run.text = "{{팁1}}"

    elif "팁 2:" in text:
        for run in runs:
            if "문화적 맥락 또는 학습 팁을 입력하세요" in run.text:
                run.text = "{{팁2}}"

    elif "팁 3:" in text:
        for run in runs:
            if "문화적 맥락 또는 학습 팁을 입력하세요" in run.text:
                run.text = "{{팁3}}"

    elif "팁 4:" in text:
        for run in runs:
            if "문화적 맥락 또는 학습 팁을 입력하세요" in run.text:
                run.text = "{{팁4}}"

    elif "자유롭게 추가 학습 내용" in text:
        for run in runs:
            run.text = ""
        if runs:
            runs[0].text = "{{메모}}"

# ── 테이블 수정 ─────────────────────────────────────────────
tables = doc.tables

# Table 0: 대화
t0 = tables[0]
for i, row in enumerate(t0.rows[1:], 1):  # 헤더 건너뜀
    set_cell_text(row.cells[0], f"{{{{대화{i}_화자}}}}")
    set_cell_text(row.cells[1], f"{{{{대화{i}_대사}}}}")

# Table 1: 핵심 표현
t1 = tables[1]
for i, row in enumerate(t1.rows[1:], 1):
    set_cell_text(row.cells[0], f"{{{{표현{i}}}}}")
    set_cell_text(row.cells[1], f"{{{{표현{i}_설명}}}}")

# Table 2: 어휘
t2 = tables[2]
for i, row in enumerate(t2.rows[1:], 1):
    set_cell_text(row.cells[0], f"{{{{단어{i}}}}}")
    set_cell_text(row.cells[1], f"{{{{단어{i}_뜻}}}}")

# Table 3: 퀴즈
t3 = tables[3]
set_cell_text(t3.rows[0].cells[1], "{{지난_정답}}")
set_cell_text(t3.rows[1].cells[1], "{{이번_퀴즈}}")
set_cell_text(t3.rows[2].cells[1], "{{퀴즈_정답}}")

doc.save(str(DST))
print(f"저장 완료: {DST}")

# 생성된 플레이스홀더 목록 출력
import re
doc2 = Document(str(DST))
found = set()
for para in doc2.paragraphs:
    for m in re.findall(r"\{\{(.+?)\}\}", para.text):
        found.add(m)
for table in doc2.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for m in re.findall(r"\{\{(.+?)\}\}", para.text):
                    found.add(m)

print(f"\n생성된 플레이스홀더 ({len(found)}개):")
for p in sorted(found):
    print(f"  {{{{ {p} }}}}")
