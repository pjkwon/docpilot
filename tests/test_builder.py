from __future__ import annotations

from pathlib import Path

import pytest

from docpilot.builder.hwpx_builder import HwpxBuilder
from docpilot.exceptions import BuilderError

SECTIONS = {"서론": "서론 내용입니다.", "결론": "결론 내용입니다."}


class TestHwpxBuilder:
    def test_build_produces_file(self, hwpx_template: Path, tmp_path: Path):
        output = tmp_path / "output.hwpx"
        builder = HwpxBuilder()
        result = builder.build(hwpx_template, SECTIONS, output)
        assert result == output
        assert output.exists()
        assert output.stat().st_size > 0

    def test_placeholders_replaced(self, hwpx_template: Path, tmp_path: Path):
        import zipfile
        output = tmp_path / "output.hwpx"
        HwpxBuilder().build(hwpx_template, SECTIONS, output)

        with zipfile.ZipFile(output, "r") as zf:
            candidates = [n for n in zf.namelist() if n.endswith("content.hml")]
            content = zf.read(candidates[0]).decode("utf-8")

        assert "서론 내용입니다." in content
        assert "{{서론}}" not in content

    def test_template_not_found_raises(self, tmp_path: Path):
        with pytest.raises(BuilderError, match="not found"):
            HwpxBuilder().build(tmp_path / "missing.hwpx", SECTIONS, tmp_path / "out.hwpx")

    def test_wrong_extension_raises(self, tmp_path: Path):
        bad = tmp_path / "template.docx"
        bad.write_bytes(b"fake")
        with pytest.raises(BuilderError, match="Expected .hwpx"):
            HwpxBuilder().build(bad, SECTIONS, tmp_path / "out.hwpx")

    def test_html_table_rowspan_colspan_becomes_merged_cells(self, hwpx_template: Path, tmp_path: Path):
        import zipfile

        html = """
        <table border="1">
          <tr><td rowspan="2">A</td><td>B</td><td>C</td></tr>
          <tr><td>D</td><td>E</td></tr>
        </table>
        """
        output = tmp_path / "output.hwpx"
        HwpxBuilder().build(hwpx_template, {"서론": html, "결론": "결론 내용입니다."}, output)

        with zipfile.ZipFile(output, "r") as zf:
            content = zf.read("Contents/content.hml").decode("utf-8")

        assert "<hp:tbl" in content
        assert 'rowSpan="2"' in content
        assert 'colSpan="1"' in content  # untouched cells still explicit 1x1

    def test_html_list_degrades_to_glyph_prefix_without_header(self, hwpx_template: Path, tmp_path: Path):
        import zipfile

        html = "<ul><li>1단계</li><ul><li style=\"list-style-type: circle;\">2단계</li></ul></ul>"
        output = tmp_path / "output.hwpx"
        HwpxBuilder().build(hwpx_template, {"서론": html, "결론": "결론 내용입니다."}, output)

        with zipfile.ZipFile(output, "r") as zf:
            content = zf.read("Contents/content.hml").decode("utf-8")

        assert "●" in content and "1단계" in content
        assert "◦" in content and "2단계" in content

    def test_html_list_creates_real_bullets_with_header(self, hwpx_template_with_header: Path, tmp_path: Path):
        import zipfile

        html = "<ul><li>1단계</li><ul><li style=\"list-style-type: circle;\">2단계</li></ul></ul>"
        output = tmp_path / "output.hwpx"
        HwpxBuilder().build(hwpx_template_with_header, {"내용": html}, output)

        with zipfile.ZipFile(output, "r") as zf:
            section = zf.read("Contents/section0.xml").decode("utf-8")
            header = zf.read("Contents/header.xml").decode("utf-8")

        assert '<hh:bullets itemCnt="2">' in header
        assert 'char="●"' in header and 'char="◦"' in header
        assert header.count('type="BULLET"') == 2
        assert "1단계" in section and "2단계" in section
        # the run text itself must not carry a literal bullet glyph — HWP draws it from header.xml
        assert "●" not in section

    def test_multi_section_all_placeholders_replaced(self, hwpx_template_multi_section: Path, tmp_path: Path):
        """section0.xml and section1.xml each carry their own {{key}} — both must be
        filled, not just the first content file found (regression: previously only
        one candidate content file was ever opened)."""
        import zipfile

        output = tmp_path / "output.hwpx"
        HwpxBuilder().build(
            hwpx_template_multi_section, {"표지": "표지 내용", "본문": "본문 내용"}, output,
        )

        with zipfile.ZipFile(output, "r") as zf:
            section0 = zf.read("Contents/section0.xml").decode("utf-8")
            section1 = zf.read("Contents/section1.xml").decode("utf-8")

        assert "표지 내용" in section0 and "{{표지}}" not in section0
        assert "본문 내용" in section1 and "{{본문}}" not in section1

    def test_multi_section_shared_bullet_para_pr_not_duplicated(
        self, hwpx_template_multi_section: Path, tmp_path: Path,
    ):
        """Two sections each rendering a same-depth bullet list must share one
        hh:paraPr in header.xml, not create a duplicate per section file."""
        import zipfile

        html = "<ul><li>항목</li></ul>"
        output = tmp_path / "output.hwpx"
        HwpxBuilder().build(hwpx_template_multi_section, {"표지": html, "본문": html}, output)

        with zipfile.ZipFile(output, "r") as zf:
            header = zf.read("Contents/header.xml").decode("utf-8")

        assert header.count('char="●"') == 1
        assert header.count('type="BULLET"') == 1

    def test_mixed_text_and_table_in_one_placeholder(self, hwpx_template: Path, tmp_path: Path):
        import zipfile

        mixed = "앞 문단\n\n| a | b |\n| --- | --- |\n| 1 | 2 |\n\n뒤 문단"
        output = tmp_path / "output.hwpx"
        HwpxBuilder().build(hwpx_template, {"서론": mixed, "결론": "결론 내용입니다."}, output)

        with zipfile.ZipFile(output, "r") as zf:
            content = zf.read("Contents/content.hml").decode("utf-8")

        assert "앞 문단" in content
        assert "<hp:tbl" in content
        assert "뒤 문단" in content
        assert content.index("앞 문단") < content.index("<hp:tbl") < content.index("뒤 문단")


class TestDocxBuilder:
    def test_build_produces_file(self, tmp_path: Path):
        docx = pytest.importorskip("docx")
        from docx import Document
        from docpilot.builder.docx_builder import DocxBuilder

        template = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("{{서론}}")
        doc.add_paragraph("{{결론}}")
        doc.save(str(template))

        output = tmp_path / "output.docx"
        DocxBuilder().build(template, SECTIONS, output)

        assert output.exists()
        result_doc = Document(str(output))
        texts = [p.text for p in result_doc.paragraphs]
        assert "서론 내용입니다." in texts
        assert "결론 내용입니다." in texts

    def test_header_and_footer_placeholders_replaced(self, tmp_path: Path):
        """Placeholders in headers/footers live in separate XML parts from the body
        (word/header1.xml, word/footer1.xml) — regression: previously only doc.paragraphs/
        doc.tables (body) were ever touched, so header/footer placeholders were silently
        left unfilled even though extraction already detects them."""
        docx = pytest.importorskip("docx")
        from docx import Document
        from docpilot.builder.docx_builder import DocxBuilder

        template = tmp_path / "template.docx"
        doc = Document()
        section = doc.sections[0]
        section.header.paragraphs[0].text = "{{작성일}}"
        section.footer.paragraphs[0].text = "{{푸터}}"
        doc.add_paragraph("{{서론}}")
        doc.save(str(template))

        output = tmp_path / "output.docx"
        DocxBuilder().build(
            template, {"작성일": "2026-07-23", "푸터": "푸터 내용", "서론": "서론 내용입니다."}, output,
        )

        result_doc = Document(str(output))
        assert [p.text for p in result_doc.sections[0].header.paragraphs] == ["2026-07-23"]
        assert [p.text for p in result_doc.sections[0].footer.paragraphs] == ["푸터 내용"]
        assert "서론 내용입니다." in [p.text for p in result_doc.paragraphs]

    def test_first_page_and_even_page_headers_replaced(self, tmp_path: Path):
        """First-page header (different_first_page_header_footer) and even-page header
        (odd_and_even_pages_header_footer) are distinct parts from the default header —
        each needs its own fill pass."""
        docx = pytest.importorskip("docx")
        from docx import Document
        from docpilot.builder.docx_builder import DocxBuilder

        template = tmp_path / "template.docx"
        doc = Document()
        section = doc.sections[0]
        section.different_first_page_header_footer = True
        section.first_page_header.paragraphs[0].text = "{{표지}}"
        doc.settings.odd_and_even_pages_header_footer = True
        section.even_page_header.paragraphs[0].text = "{{짝수헤더}}"
        doc.add_paragraph("{{서론}}")
        doc.save(str(template))

        output = tmp_path / "output.docx"
        DocxBuilder().build(
            template,
            {"표지": "표지 헤더", "짝수헤더": "짝수 헤더 내용", "서론": "서론 내용입니다."},
            output,
        )

        result_doc = Document(str(output))
        s = result_doc.sections[0]
        assert [p.text for p in s.first_page_header.paragraphs] == ["표지 헤더"]
        assert [p.text for p in s.even_page_header.paragraphs] == ["짝수 헤더 내용"]

    def test_linked_header_not_double_processed(self, tmp_path: Path):
        """A second section whose header is_linked_to_previous shares section 1's
        underlying header part — must not be visited as a separate fill pass (which
        would be redundant, and multi-line values would otherwise get cloned twice)."""
        docx = pytest.importorskip("docx")
        from docx.enum.section import WD_SECTION
        from docx import Document
        from docpilot.builder.docx_builder import DocxBuilder

        template = tmp_path / "template.docx"
        doc = Document()
        doc.sections[0].header.paragraphs[0].text = "{{작성일}}\n둘째줄"
        doc.add_section(WD_SECTION.NEW_PAGE)  # section 2's header defaults to linked
        doc.save(str(template))

        output = tmp_path / "output.docx"
        DocxBuilder().build(template, {"작성일": "2026-07-23"}, output)

        result_doc = Document(str(output))
        assert result_doc.sections[1].header.is_linked_to_previous
        texts = [p.text for p in result_doc.sections[0].header.paragraphs]
        assert texts == ["2026-07-23", "둘째줄"]
