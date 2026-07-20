"""
data/ 폴더 템플릿 + 전사 파일로 문서 생성 E2E 스크립트

실행:
    uv run python scripts/test_docgen.py
    uv run python scripts/test_docgen.py --preset 학회
    uv run python scripts/test_docgen.py --convert
    uv run python scripts/test_docgen.py --dry-run
    uv run python scripts/test_docgen.py --template data/my.docx --source data/my.txt --output output/result.docx
"""
from __future__ import annotations

import argparse
import io
import os
import shutil
import sys
import tempfile
import time
from pathlib import Path

# Windows 터미널 CP949 인코딩 오류 방지
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

ROOT = Path(__file__).parent.parent
DATA = ROOT / "data"
OUTPUT = ROOT / "output"

# ── 프리셋 정의 ────────────────────────────────────────────────────────────────
PRESETS: dict[str, dict] = {
    "ebs": {
        "template": DATA / "EBS_영어학습_정리노트_v2.docx",
        "source":   DATA / "EBS_BEGINNER_Script_20260527.txt",
        "output":   OUTPUT / "EBS_영어학습_정리노트_출력.docx",
        "convert_keys": [],  # 이미 번호 그룹(단어1, 단어2...)으로 구성됨
    },
    "학회": {
        "template": DATA / "학회참석보고서_템플릿_1.docx",
        "source":   DATA / "EBS_BEGINNER_Script_20260527.txt",
        "output":   OUTPUT / "학회참석보고서_출력.docx",
        "convert_keys": [
            "주요 세션별 발표 내용 요약",
            "기조 강연 주요 내용 및 발표자",
            "학회에서 파악한 최신 기술 동향 및 시사점",
            "주목할 만한 논문 및 포스터 발표 내용",
        ],
    },
}


def _extract_placeholders(template: Path) -> list[str]:
    """Full placeholder strings ({{key}}, {{?key?}} 등) 그대로 반환."""
    import re
    _RE = re.compile(r"\{\{\??[^}?]+\??\}\}")
    placeholders: list[str] = []
    seen: set[str] = set()

    if template.suffix.lower() == ".docx":
        import docx
        doc = docx.Document(str(template))
        paras = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paras.extend(cell.paragraphs)
        for para in paras:
            for m in _RE.finditer(para.text):
                s = m.group(0)
                if s not in seen:
                    seen.add(s)
                    placeholders.append(s)
    else:
        import zipfile
        with zipfile.ZipFile(template) as zf:
            candidates = [n for n in zf.namelist() if "content.hml" in n or "section0.xml" in n]
            for name in candidates:
                text = zf.read(name).decode("utf-8", errors="ignore")
                for m in _RE.finditer(text):
                    s = m.group(0)
                    if s not in seen:
                        seen.add(s)
                        placeholders.append(s)
    return placeholders


def _print_template_info(template: Path, keys: list[str]) -> None:
    print(f"템플릿    : {template.name}")
    print(f"플레이스홀더 ({len(keys)}개):")
    for k in keys:
        print(f"  {k}")
    print()


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--preset", choices=list(PRESETS.keys()), default="ebs", help="테스트 프리셋 (기본: ebs)")
    parser.add_argument("--template", type=Path, help="템플릿 파일 경로 (.docx / .hwpx), 프리셋 무시")
    parser.add_argument("--source", type=Path, help="소스 파일 경로 (.txt / .md 등), 프리셋 무시")
    parser.add_argument("--output", type=Path, help="출력 파일 경로, 프리셋 무시")
    parser.add_argument("--convert", action="store_true", help="가변 플레이스홀더({{?key?}})로 변환 후 생성")
    parser.add_argument("--convert-keys", nargs="+", metavar="KEY", help="변환할 플레이스홀더 키 직접 지정 (space 포함 시 따옴표 사용)")
    parser.add_argument("--dry-run", action="store_true", help="비용 추정만 (API 호출 없음)")
    return parser.parse_args()


def main() -> None:
    args = parse_args()

    preset = PRESETS[args.preset]
    template_path: Path = args.template or preset["template"]
    source_path: Path   = args.source   or preset["source"]
    output_path: Path   = args.output   or preset["output"]
    convert_keys: list[str] = args.convert_keys or preset.get("convert_keys", [])

    OUTPUT.mkdir(exist_ok=True)

    # ── 검증 ──────────────────────────────────────────────────────────────────
    if not template_path.exists():
        print(f"[오류] 템플릿 파일 없음: {template_path}")
        sys.exit(1)
    if not source_path.exists():
        print(f"[오류] 소스 파일 없음: {source_path}")
        sys.exit(1)

    # ── 플레이스홀더 표시 ──────────────────────────────────────────────────────
    print("=" * 60)
    keys = _extract_placeholders(template_path)
    _print_template_info(template_path, keys)

    # ── 가변 변환 ──────────────────────────────────────────────────────────────
    active_template = template_path
    if args.convert and convert_keys:
        from docpilot import convert_to_list_placeholder
        converted_out = OUTPUT / (template_path.stem + "_converted" + template_path.suffix)
        print(f"[변환] 다음 키를 {{{{?key?}}}}으로 변환:")
        for k in convert_keys:
            print(f"  · {k}")
        active_template = convert_to_list_placeholder(template_path, convert_keys, output=converted_out)
        print(f"  → 저장: {active_template}\n")

        # 변환 후 플레이스홀더 다시 출력
        converted_keys = _extract_placeholders(active_template)
        print(f"변환 후 플레이스홀더 ({len(converted_keys)}개):")
        for k in converted_keys:
            print(f"  {k}")
        print()
    elif args.convert and not convert_keys:
        print("[알림] 이 프리셋에 변환 대상 키가 없습니다. --convert-keys 로 직접 지정하세요.\n")

    # ── DocPilot 초기화 ────────────────────────────────────────────────────────
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("[오류] ANTHROPIC_API_KEY 환경변수가 필요합니다.")
        sys.exit(1)

    import docpilot
    pilot = docpilot.DocPilot(llm="claude", api_key=api_key)

    # 소스 파일을 임시 폴더에 복사 (generate()는 data_folder 단위로 처리)
    tmp_dir = tempfile.mkdtemp(prefix="docpilot_")
    try:
        shutil.copy2(source_path, tmp_dir)
        print(f"소스      : {source_path.name}")
        print(f"출력      : {output_path}")
        print()

        # ── 비용 추정 ─────────────────────────────────────────────────────────
        if args.dry_run:
            print("── 비용 추정 (quick 모드) ──────────────────────────────────────")
            report = pilot.estimate_cost(tmp_dir, active_template, quick=True)
            print(report)
            return

        # ── 문서 생성 ─────────────────────────────────────────────────────────
        print("── 문서 생성 ────────────────────────────────────────────────────")
        t0 = time.perf_counter()
        try:
            result = pilot.generate(
                data_folder=tmp_dir,
                template=active_template,
                output=output_path,
                reindex=True,
            )
        except Exception as e:
            print(f"\n[실패] {e}")
            sys.exit(1)

        elapsed = time.perf_counter() - t0
        print(f"\n── 완료 ─────────────────────────────────────────────────────────")
        print(f"출력 파일 : {output_path}")
        print(f"총 소요   : {elapsed:.1f}s")
        print(f"모델      : {result.model}")
        print(f"토큰      : 입력 {result.input_tokens:,} / 출력 {result.output_tokens:,}")

    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


if __name__ == "__main__":
    main()
