"""
HWPX → DOCX 변환 테스트 (한컴오피스 COM 자동화)

실행:
    uv run python scripts/test_hwpx_to_docx.py
"""
from __future__ import annotations

import sys
import time
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

ROOT   = Path(__file__).parent.parent
SOURCE = ROOT / "data" / "(양식2) 연구개발계획서(국문).hwpx"
OUTPUT = ROOT / "output" / "(양식2) 연구개발계획서(국문).docx"


def main() -> None:
    if not SOURCE.exists():
        print(f"[오류] 소스 파일 없음: {SOURCE}")
        sys.exit(1)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    from docpilot.builder.hwp_convert import convert_to_docx
    from docpilot.exceptions import ConversionError

    print(f"입력   : {SOURCE.name}")
    print(f"출력   : {OUTPUT}")
    print()

    t0 = time.perf_counter()
    try:
        result = convert_to_docx(SOURCE, OUTPUT)
    except ConversionError as e:
        print(f"\n[실패] {e}")
        sys.exit(1)
    elapsed = time.perf_counter() - t0

    print(f"완료   : {result}")
    print(f"소요   : {elapsed:.1f}s")

    import docx
    doc = docx.Document(str(result))
    print(f"단락 수 : {len(doc.paragraphs)}")
    print(f"표 개수 : {len(doc.tables)}")


if __name__ == "__main__":
    main()
