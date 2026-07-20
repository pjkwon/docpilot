"""
DOCX 템플릿 구조를 참고해 HWPX로 생성

HwpxDynamicBuilder: HWPX 템플릿 없이도 LLM이 구조를 만들어 HWPX 출력 가능.
DOCX 플레이스홀더 목록을 지침으로 넘겨 같은 섹션 구조를 재현한다.

실행:
    uv run python scripts/gen_학회보고서_hwpx_from_docx.py
"""
from __future__ import annotations

import os
import re
import sys
import time
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

ROOT     = Path(__file__).parent.parent
TEMPLATE = ROOT / "data" / "학회참석보고서_템플릿_1.docx"
SOURCE   = ROOT / "data" / "0617_제주학회_발표정리.md"
OUTPUT   = ROOT / "output" / "0617_제주학회_참석보고서_from_docx.hwpx"
HEADER   = ROOT / "docpilot" / "templates" / "base" / "Contents" / "header.xml"


def extract_docx_placeholders(path: Path) -> list[str]:
    import docx
    _RE = re.compile(r"\{\{\??([^}?]+)\??\}\}")
    seen: dict[str, None] = {}
    doc = docx.Document(str(path))
    paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)
    for para in paras:
        for m in _RE.finditer(para.text):
            seen[m.group(1)] = None
    return list(seen)


def main() -> None:
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("[오류] ANTHROPIC_API_KEY 환경변수가 필요합니다.")
        sys.exit(1)

    for p in (TEMPLATE, SOURCE, HEADER):
        if not p.exists():
            print(f"[오류] 파일 없음: {p}")
            sys.exit(1)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    # DOCX 플레이스홀더 목록 추출 → LLM 지침으로 활용
    keys = extract_docx_placeholders(TEMPLATE)
    print(f"DOCX 플레이스홀더 ({len(keys)}개):")
    for k in keys:
        print(f"  · {k}")
    print()

    sections_hint = "\n".join(f"- {k}" for k in keys)
    instructions = (
        "학회 출장 참석 보고서를 공식 문서 형식으로 작성하세요. "
        "경어체(~하였습니다, ~입니다)를 사용하세요.\n\n"
        "반드시 아래 항목들을 섹션으로 구성하세요:\n"
        f"{sections_hint}"
    )

    content = SOURCE.read_text(encoding="utf-8")

    from docpilot.builder.hwpx_dynamic_builder import HwpxDynamicBuilder
    from docpilot.mapping.claude import ClaudeMapper

    mapper = ClaudeMapper(api_key=api_key, max_tokens=16000)
    builder = HwpxDynamicBuilder(mapper=mapper)

    print(f"소스   : {SOURCE.name}")
    print(f"출력   : {OUTPUT}")
    print()

    t0 = time.perf_counter()
    try:
        result = builder.build(
            content=content,
            header_xml=HEADER,
            output=OUTPUT,
            instructions=instructions,
        )
    except Exception as e:
        print(f"\n[실패] {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

    elapsed = time.perf_counter() - t0
    print(f"\n완료   : {result}")
    print(f"소요   : {elapsed:.1f}s")


if __name__ == "__main__":
    main()
