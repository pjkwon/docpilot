from __future__ import annotations

import copy
import re
from pathlib import Path

from docpilot.builder.base import (
    BaseBuilder,
    PLACEHOLDER_RE,
    OPTIONAL_PLACEHOLDER_RE,
    _NUMBERED_GROUP_RE,
)
from docpilot.exceptions import BuilderError

# DOCX XML namespace
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W_T = f"{{{_W_NS}}}t"
_W_TR = f"{{{_W_NS}}}tr"
_W_TC = f"{{{_W_NS}}}tc"


def _iter_containers(doc):
    """Yield the document body plus every header/footer that has its own
    definition — default, first-page, and even-page, for every section.

    A header/footer with no definition of its own (``is_linked_to_previous``)
    shares the prior section's part; skipping it avoids processing the same
    underlying XML twice and avoids python-docx silently creating a blank
    definition just because we touched ``.paragraphs`` on it.
    """
    yield doc
    for section in doc.sections:
        for hdr_ftr in (
            section.header, section.footer,
            section.first_page_header, section.first_page_footer,
            section.even_page_header, section.even_page_footer,
        ):
            if not hdr_ftr.is_linked_to_previous:
                yield hdr_ftr


class DocxBuilder(BaseBuilder):
    def build(
        self,
        template: str | Path,
        sections: dict[str, str | list[str]],
        output: str | Path,
    ) -> Path:
        template, output = self._validate_paths(template, output)

        if template.suffix.lower() != ".docx":
            raise BuilderError(f"Expected .docx template, got '{template.suffix}'")

        try:
            import docx
        except ImportError as e:
            raise BuilderError("python-docx is required: pip install python-docx") from e

        try:
            doc = docx.Document(str(template))
        except Exception as e:
            raise BuilderError("Failed to open DOCX template", detail=str(e)) from e

        # Body + every header/footer part get the same three-phase treatment —
        # placeholders can live in any of them, not just the body.
        for container in _iter_containers(doc):
            # ── Phase 1: Table group expansion / contraction ──────────────────
            for table in container.tables:
                _process_table_groups(table, sections)

            # ── Phase 2: Fill all paragraphs and table cells ───────────────────
            for para in list(container.paragraphs):
                _replace_in_paragraph(para, sections)

            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in list(cell.paragraphs):
                            _replace_in_paragraph(para, sections)

            # ── Phase 3: Remove table rows where every cell is now empty ───────
            for table in container.tables:
                _remove_empty_optional_rows(table, sections)

        try:
            doc.save(str(output))
        except Exception as e:
            raise BuilderError("Failed to save DOCX", detail=str(e)) from e

        return output


# ── Table group processing ────────────────────────────────────────────────────

def _process_table_groups(table, sections: dict) -> None:
    """Expand or contract rows for numbered optional list groups."""
    tbl_xml = table._tbl

    # Find rows that belong to each numbered group
    # group_row_map: base → [(n, tr_element), ...]
    group_row_map: dict[str, list[tuple[int, object]]] = {}
    group_tr_by_n: dict[str, dict[int, object]] = {}  # base → {n: tr}

    all_tr = list(tbl_xml.findall(_W_TR))

    for tr in all_tr:
        seen_in_row: set[str] = set()
        for t_elem in tr.iter(_W_T):
            if not t_elem.text:
                continue
            for m in OPTIONAL_PLACEHOLDER_RE.finditer(t_elem.text):
                key = m.group(1)
                ng = _NUMBERED_GROUP_RE.match(key)
                if ng:
                    base = ng.group(1) + ng.group(3)
                    n = int(ng.group(2))
                    if isinstance(sections.get(base), list) and (base, n) not in seen_in_row:
                        seen_in_row.add((base, n))
                        group_row_map.setdefault(base, []).append((n, tr))
                        group_tr_by_n.setdefault(base, {})[n] = tr

    if not group_row_map:
        return

    # Each group that shares the same set of rows should be expanded together.
    # We process by unique row-set (multiple bases may share the same rows — e.g. 단어, 단어_뜻).
    processed_rows: set[id] = set()

    # Sort groups by their first occurrence to avoid double-processing
    for base, row_list in group_row_map.items():
        first_tr_id = id(row_list[0][1])
        if first_tr_id in processed_rows:
            continue

        # Find all bases that share ANY row with this base
        sibling_rows_set = {id(tr) for _, tr in row_list}
        sibling_bases = [
            b for b, rl in group_row_map.items()
            if any(id(tr) in sibling_rows_set for _, tr in rl)
        ]

        # Template rows sorted by n
        sorted_rows = sorted(row_list, key=lambda x: x[0])
        template_ns = [n for n, _ in sorted_rows]
        template_trs = [tr for _, tr in sorted_rows]
        template_count = len(template_trs)
        last_tr = template_trs[-1]
        last_n = template_ns[-1]

        # Target count = length of the list value (use first sibling base)
        target_count = max(
            len(sections[b]) for b in sibling_bases
            if isinstance(sections.get(b), list)
        )

        if target_count > template_count:
            # Clone last template row, renumbering placeholders
            insert_after = last_tr
            for new_n in range(last_n + 1, last_n + (target_count - template_count) + 1):
                new_tr = copy.deepcopy(last_tr)
                _renumber_placeholders_in_tr(new_tr, last_n, new_n)
                insert_after.addnext(new_tr)
                insert_after = new_tr

        elif target_count < template_count:
            # Remove excess rows (from last n down to target_count+1)
            for _, tr in reversed(sorted_rows[target_count:]):
                tbl_xml.remove(tr)

        for _, tr in row_list:
            processed_rows.add(id(tr))


def _renumber_placeholders_in_tr(tr, old_n: int, new_n: int) -> None:
    """In a cloned <w:tr>, replace {{?keyOLD_N}} → {{?keyNEW_N}} in all text nodes."""
    for t_elem in tr.iter(_W_T):
        if t_elem.text and "{{?" in t_elem.text:
            def replace_n(m, _old=old_n, _new=new_n):
                key = m.group(1)
                ng = _NUMBERED_GROUP_RE.match(key)
                if ng and int(ng.group(2)) == _old:
                    return f"{{{{?{ng.group(1)}{_new}{ng.group(3)}}}}}"
                return m.group(0)
            t_elem.text = OPTIONAL_PLACEHOLDER_RE.sub(replace_n, t_elem.text)


def _remove_empty_optional_rows(table, sections: dict) -> None:
    """Remove table rows where every cell's optional placeholder was left empty."""
    tbl_xml = table._tbl
    to_remove = []

    for tr in tbl_xml.findall(_W_TR):
        all_texts = []
        has_optional = False
        for t_elem in tr.iter(_W_T):
            txt = t_elem.text or ""
            all_texts.append(txt)
            # Check if any optional placeholder remains OR the cell was cleared to ""
            for m in OPTIONAL_PLACEHOLDER_RE.finditer(txt):
                has_optional = True

        # If the row still contains unfilled optional placeholders → all-empty row
        if has_optional:
            to_remove.append(tr)
            continue

        # Also remove rows where all visible text is empty (optional section was "" )
        combined = "".join(all_texts).strip()
        # Only auto-remove if the row had at least one numbered-group optional cell
        # (checked by presence of list-type values in sections for keys in this row)
        row_has_list_key = False
        for t_elem in tr.iter(_W_T):
            txt = t_elem.text or ""
            # After fill, list-group cells look like normal text; we rely on the
            # unfilled placeholder check above for removal.
        if not combined and row_has_list_key:
            to_remove.append(tr)

    for tr in to_remove:
        try:
            tbl_xml.remove(tr)
        except Exception:
            pass


# ── Paragraph replacement ─────────────────────────────────────────────────────

def _replace_in_paragraph(para, sections: dict[str, str | list[str]]) -> None:
    full_text = "".join(run.text for run in para.runs)
    if not PLACEHOLDER_RE.search(full_text):
        return

    def replacer(m):
        key = m.group(1)
        # Is this a numbered group item? e.g. "단어1", "단어1_뜻"
        ng = _NUMBERED_GROUP_RE.match(key)
        if ng:
            base = ng.group(1) + ng.group(3)
            n = int(ng.group(2))
            val = sections.get(base)
            if isinstance(val, list):
                idx = n - 1
                return val[idx] if idx < len(val) else ""
        # Regular or optional single
        val = sections.get(key)
        if val is None:
            return m.group(0)  # not in sections → leave as-is
        if isinstance(val, list):
            return "\n".join(val) if val else ""
        return val

    replaced = PLACEHOLDER_RE.sub(replacer, full_text)
    if replaced == full_text:
        return

    lines = replaced.split("\n")

    # Write first line back into the original paragraph (preserves run formatting)
    if para.runs:
        para.runs[0].text = lines[0]
        for run in para.runs[1:]:
            run.text = ""

    if len(lines) == 1:
        return

    # Multi-line: clone paragraph XML for each subsequent line
    from docx.oxml.ns import qn

    p_elem = para._element
    parent = p_elem.getparent()
    if parent is None:
        if para.runs:
            para.runs[0].text = " ".join(lines)
        return

    idx = list(parent).index(p_elem)
    for i, line in enumerate(lines[1:], 1):
        new_p = copy.deepcopy(p_elem)
        runs = new_p.findall(qn("w:r"))
        if runs:
            t_list = runs[0].findall(qn("w:t"))
            if t_list:
                t_list[0].text = line
                for extra_t in t_list[1:]:
                    runs[0].remove(extra_t)
            for extra_r in runs[1:]:
                new_p.remove(extra_r)
        parent.insert(idx + i, new_p)
