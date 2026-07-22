from __future__ import annotations

import atexit
import os
import re
from importlib.metadata import version, PackageNotFoundError
from typing import TYPE_CHECKING

try:
    __version__ = version("smart-docgen")
except PackageNotFoundError:
    __version__ = "0.0.0+unknown"
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path

from dotenv import load_dotenv

from docpilot.exceptions import BuilderError, DocPilotError, MappingError

if TYPE_CHECKING:
    from docpilot.ingestion.models import IngestedDocument

load_dotenv()

_PLACEHOLDER_RE = re.compile(r"\{\{\??(.+?)\??\}\}")      # {{key}}, {{?key}}, {{?key?}}
_OPTIONAL_PLACEHOLDER_RE = re.compile(r"\{\{\?(.+?)\}\}")  # only {{?key}} and {{?key?}}
_DYNAMIC_LIST_RE = re.compile(r"\{\{\?([^}?]+)\?\}\}")      # only {{?key?}}
_NUMBERED_GROUP_RE = re.compile(r"^(.+?)(\d+)(.*)$")        # "단어1_뜻" → ("단어","1","_뜻")

@dataclass
class GenerateResult:
    path: Path
    model: str
    input_tokens: int
    output_tokens: int
    elapsed_seconds: float

    @property
    def total_tokens(self) -> int:
        return self.input_tokens + self.output_tokens

    def __str__(self) -> str:
        return str(self.path)

    def __fspath__(self) -> str:
        return str(self.path)


_MODEL_PRICING: dict[str, tuple[float, float]] = {
    "claude-opus-4-8":   (5.00, 25.00),
    "claude-opus-4-7":   (5.00, 25.00),
    "claude-opus-4-6":   (5.00, 25.00),
    "claude-sonnet-4-6": (3.00, 15.00),
    "claude-haiku-4-5":  (1.00,  5.00),
}
_EST_OUTPUT_TOKENS_PER_SECTION = 500

# Maps file extension → required extra (None = included in base install)
_EXT_EXTRAS: dict[str, str | None] = {
    ".txt":  None,
    ".md":   None,
    ".rst":  None,
    ".csv":  None,
    ".hwpx": None,
    ".hwp":  "hwp",
    ".pdf":  "pdf",
    ".docx": "docx",
}


def suggest_extras(folder: str | Path) -> dict:
    """
    Scan a data folder and suggest which docpilot extras to install.

    Returns a dict with:
      found          — {extension: file_count} for all files found
      unsupported    — {extension: file_count} for files docpilot cannot process
      required_extras — list of extras needed (e.g. ["pdf", "docx"])
      install_command — ready-to-run pip command, or None if nothing extra needed
    """
    folder = Path(folder)
    if not folder.is_dir():
        raise DocPilotError("Not a directory", detail=str(folder))

    found: dict[str, int] = {}
    for file in folder.rglob("*"):
        if not file.is_file():
            continue
        ext = file.suffix.lower()
        found[ext] = found.get(ext, 0) + 1

    unsupported: dict[str, int] = {
        ext: count for ext, count in found.items() if ext not in _EXT_EXTRAS
    }

    required: list[str] = sorted({
        extra
        for ext in found
        if ext in _EXT_EXTRAS and (extra := _EXT_EXTRAS[ext]) is not None
    })

    install_command = (
        f'pip install "smart-docgen[{",".join(required)}]"' if required else None
    )

    return {
        "found": found,
        "unsupported": unsupported,
        "required_extras": required,
        "install_command": install_command,
    }
_BUILTIN_TEMPLATES = Path(__file__).parent / "templates"


def _get_builtin_metadata() -> dict[str, str]:
    from docpilot.mapping.sidecar import load_sidecar
    result = {}
    for d in sorted(_BUILTIN_TEMPLATES.iterdir()):
        if not d.is_dir() or d.name == "base":
            continue
        sidecar = load_sidecar(d)
        if sidecar and sidecar.description:
            result[d.name] = sidecar.description
        else:
            result[d.name] = d.name
    return result

_ASSEMBLED_CACHE: dict[str, Path] = {}


def _assemble_builtin_hwpx(name: str) -> Path:
    """Assemble a built-in HWPX template from base + per-template XML sources."""
    if name in _ASSEMBLED_CACHE:
        return _ASSEMBLED_CACHE[name]

    base_dir = _BUILTIN_TEMPLATES / "base"
    overlay_dir = _BUILTIN_TEMPLATES / name

    if not base_dir.exists() or not overlay_dir.exists():
        raise DocPilotError(
            f"Built-in template '{name}' not found",
            detail=f"Expected source at {overlay_dir}",
        )

    tmp = tempfile.NamedTemporaryFile(suffix=f"_{name}.hwpx", delete=False)
    tmp.close()
    tmp_path = Path(tmp.name)

    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zf:
        mimetype = base_dir / "mimetype"
        if mimetype.exists():
            zf.write(mimetype, "mimetype", compress_type=zipfile.ZIP_STORED)

        for file in sorted(base_dir.rglob("*")):
            if not file.is_file() or file.name == "mimetype":
                continue
            arcname = file.relative_to(base_dir)
            # overlay_dir files are named header.xml / section0.xml — placed in Contents/
            override = overlay_dir / file.name
            source = override if override.exists() else file
            zf.write(source, arcname.as_posix())

    _ASSEMBLED_CACHE[name] = tmp_path
    atexit.register(lambda p=tmp_path: p.unlink(missing_ok=True))
    return tmp_path


def _m(model: str | None) -> dict:
    return {"model": model} if model else {}


def _b(base_url: str | None) -> dict:
    return {"base_url": base_url} if base_url else {}


def _ingest_instructions_doc(path: Path) -> str:
    """지침 문서를 읽어 텍스트로 반환한다. 지원하지 않는 형식이면 빈 문자열 반환."""
    from docpilot.ingestion import text as text_ing
    from docpilot.ingestion import hwpx as hwpx_ing

    ext = path.suffix.lower()

    try:
        if ext == ".hwpx":
            return hwpx_ing.ingest(path).content
        if ext in text_ing.SUPPORTED_EXTENSIONS:
            return text_ing.ingest(path).content
        if ext == ".pdf":
            from docpilot.ingestion import pdf as pdf_ing
            return pdf_ing.ingest(path).content
        if ext == ".docx":
            from docpilot.ingestion import docx as docx_ing
            return docx_ing.ingest(path).content
    except Exception:
        pass
    return ""


def _validate_hwpx(path: Path) -> None:
    """생성된 HWPX 파일의 구조 무결성을 비차단 방식으로 검사한다."""
    import logging
    import warnings

    _REQUIRED = [
        "mimetype",
        "Contents/content.hpf",
        "Contents/header.xml",
    ]

    try:
        with zipfile.ZipFile(path, "r") as zf:
            names = zf.namelist()
            missing = [f for f in _REQUIRED if f not in names]
            if missing:
                warnings.warn(
                    f"HWPX 검증 경고 — 필수 파일 누락: {', '.join(missing)}",
                    stacklevel=4,
                )

            try:
                from lxml import etree
                for name in names:
                    if name.endswith(".xml") or name.endswith(".hpf"):
                        try:
                            etree.fromstring(zf.read(name))
                        except Exception as xml_err:
                            warnings.warn(
                                f"HWPX 검증 경고 — XML 오류 in {name}: {xml_err}",
                                stacklevel=4,
                            )
            except ImportError:
                pass
    except Exception as exc:
        logging.getLogger(__name__).debug("HWPX 검증 실패 (무시됨): %s", exc)


def _estimate_tokens_from_folder(folder: Path, n_sections: int) -> int:
    """Rough token estimate from data folder file sizes. ~0.4 tokens/byte for Korean text."""
    _TEXT_EXTS = {".txt", ".md", ".rst", ".csv", ".hwpx", ".docx"}
    _BINARY_EXTS = {".pdf"}
    total_bytes = 0
    for f in folder.rglob("*"):
        if not f.is_file():
            continue
        ext = f.suffix.lower()
        if ext in _TEXT_EXTS:
            total_bytes += f.stat().st_size
        elif ext in _BINARY_EXTS:
            total_bytes += f.stat().st_size * 3  # OCR/extracted text is larger than raw binary
    tokens_per_byte = 0.4
    raw_tokens = int(total_bytes * tokens_per_byte)
    # RAG retrieves a subset, not the full corpus — cap at ~4,000 tokens per section
    return min(raw_tokens, n_sections * 4000)


_DEFAULT_MAX_INPUT_TOKENS_WARN = 50_000


def _estimate_tokens_raw(text: str) -> int:
    """Rough token estimate for text that goes into the prompt in full — no RAG retrieval
    subset, so (unlike _estimate_tokens_from_folder) no per-section cap applies."""
    return int(len(text.encode("utf-8")) * 0.4)


def _resolve_instructions(
    extra_instructions: str | None,
    sidecar_instructions: str,
    instructions_doc: str | Path | None,
) -> str | None:
    """Merge sidecar instructions + instructions_doc content into extra_instructions.

    Priority: instructions_doc is prepended last (highest priority — read in full,
    unlike sidecar/extra which are just strings), sidecar next, caller's extra_instructions last.
    """
    if sidecar_instructions:
        extra_instructions = (
            f"{sidecar_instructions}\n\n{extra_instructions}"
            if extra_instructions
            else sidecar_instructions
        )

    if instructions_doc is not None:
        doc_text = _ingest_instructions_doc(Path(instructions_doc))
        if doc_text:
            header = f"[지침 문서: {Path(instructions_doc).name}]\n{doc_text}"
            extra_instructions = (
                f"{header}\n\n{extra_instructions}" if extra_instructions
                else header
            )

    return extra_instructions


def _resolve_content(content: str | list[str | Path] | list[IngestedDocument]) -> str:
    """Resolve the content union type shared by *_from_content() methods into a single string."""
    from docpilot.ingestion.models import IngestedDocument
    from docpilot.mapping.base import merge_documents

    if isinstance(content, str):
        return content
    if content and isinstance(content[0], IngestedDocument):
        return merge_documents(content)

    from docpilot.ingestion import ingest_paths
    docs = ingest_paths(content)
    if not docs:
        raise DocPilotError("content에서 ingest 가능한 파일을 찾지 못했습니다.", detail=str(content))
    return merge_documents(docs)


def _check_content_size(content_str: str, max_input_tokens: int | None) -> None:
    """Warn or raise on a rough byte-based token estimate — *_from_content() methods send
    content in full with no RAG top_k to bound its size automatically."""
    estimated_tokens = _estimate_tokens_raw(content_str)
    if max_input_tokens is not None:
        if estimated_tokens > max_input_tokens:
            raise MappingError(
                f"콘텐츠 예상 입력 토큰 수({estimated_tokens:,})가 "
                f"max_input_tokens({max_input_tokens:,})를 초과합니다.",
                detail="content 크기를 줄이거나 max_input_tokens를 늘리세요.",
            )
    elif estimated_tokens > _DEFAULT_MAX_INPUT_TOKENS_WARN:
        import warnings
        warnings.warn(
            f"콘텐츠 예상 입력 토큰 수가 {estimated_tokens:,}로 큽니다 "
            f"(경고 기준: {_DEFAULT_MAX_INPUT_TOKENS_WARN:,}). RAG와 달리 이 경로는 "
            "content를 통째로 프롬프트에 넣으므로 컨텍스트 초과·비용 증가 위험이 있습니다. "
            "content를 줄이거나 max_input_tokens로 명시적 상한을 거세요.",
            stacklevel=2,
        )


def _extract_placeholders(template_path: Path) -> list[str]:
    """Extract {{section}} placeholder names from a template file."""
    suffix = template_path.suffix.lower()

    if suffix == ".hwpx":
        with zipfile.ZipFile(template_path, "r") as zf:
            names = zf.namelist()
            candidates = [n for n in names if n.endswith("content.hml")]
            if not candidates:
                candidates = [n for n in names if n.endswith("section0.xml")]
            if not candidates:
                return []
            text = "".join(
                zf.read(c).decode("utf-8", errors="ignore") for c in candidates
            )
    elif suffix == ".docx":
        with zipfile.ZipFile(template_path, "r") as zf:
            text = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    elif suffix == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(template_path) as pdf:
                text = "\n".join(p.extract_text() or "" for p in pdf.pages)
        except Exception:
            return []
    else:
        return []

    seen: dict[str, None] = {}
    for match in _PLACEHOLDER_RE.finditer(text):
        seen[match.group(1)] = None
    return list(seen)


def _extract_placeholder_sections(template_path: Path) -> list:
    """
    Extract placeholders and return TemplateSection list with optional/list flags.

    - {{key}}  → required TemplateSection
    - {{?key}} → optional TemplateSection (optional=True)
    - {{?단어1}}, {{?단어2}}, ... → collapsed to TemplateSection(name="단어", is_list=True)
    """
    from docpilot.mapping.base import TemplateSection

    suffix = template_path.suffix.lower()

    if suffix == ".hwpx":
        with zipfile.ZipFile(template_path, "r") as zf:
            names = zf.namelist()
            candidates = [n for n in names if n.endswith("content.hml")]
            if not candidates:
                candidates = [n for n in names if n.endswith("section0.xml")]
            if not candidates:
                return []
            text = "".join(
                zf.read(c).decode("utf-8", errors="ignore") for c in candidates
            )
    elif suffix == ".docx":
        with zipfile.ZipFile(template_path, "r") as zf:
            text = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    elif suffix == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(template_path) as pdf:
                text = "\n".join(p.extract_text() or "" for p in pdf.pages)
        except Exception:
            return []
    else:
        return []

    # Collect all keys in first-seen order
    all_keys: list[str] = []
    seen_keys: set[str] = set()
    optional_keys: set[str] = set()
    dynamic_list_keys: set[str] = set()  # {{?key?}}

    for m in _DYNAMIC_LIST_RE.finditer(text):
        dynamic_list_keys.add(m.group(1))

    for m in _OPTIONAL_PLACEHOLDER_RE.finditer(text):
        optional_keys.add(m.group(1))

    for m in _PLACEHOLDER_RE.finditer(text):
        key = m.group(1)
        if key not in seen_keys:
            all_keys.append(key)
            seen_keys.add(key)

    # Detect numbered optional groups: {{?단어1}}, {{?단어2}}, ... → base "단어"
    numbered_groups: dict[str, int] = {}  # base → max N
    numbered_first_key: dict[str, str] = {}  # base → first raw key (for style_hint lookup)
    for key in all_keys:
        if key in optional_keys and key not in dynamic_list_keys:
            m = _NUMBERED_GROUP_RE.match(key)
            if m:
                base = m.group(1) + m.group(3)
                n = int(m.group(2))
                if base not in numbered_groups:
                    numbered_groups[base] = n
                    numbered_first_key[base] = key
                else:
                    numbered_groups[base] = max(numbered_groups[base], n)

    # Build TemplateSection list, collapsing numbered optional groups
    result: list[TemplateSection] = []
    emitted: set[str] = set()

    for key in all_keys:
        if key in dynamic_list_keys:
            # {{?key?}} → dynamic list (count determined by LLM from source data)
            if key not in emitted:
                emitted.add(key)
                result.append(TemplateSection(name=key, is_list=True, group_max=0))
        elif key in optional_keys:
            m = _NUMBERED_GROUP_RE.match(key)
            if m:
                base = m.group(1) + m.group(3)
                if base not in emitted:
                    emitted.add(base)
                    result.append(TemplateSection(
                        name=base,
                        optional=True,
                        is_list=True,
                        group_max=numbered_groups[base],
                    ))
                continue
            if key not in emitted:
                emitted.add(key)
                result.append(TemplateSection(name=key, optional=True))
        else:
            if key not in emitted:
                emitted.add(key)
                result.append(TemplateSection(name=key))

    return result


def convert_to_list_placeholder(
    path: str | Path,
    keys: list[str] | str,
    output: str | Path | None = None,
) -> Path:
    """
    Convert {{key}} placeholders to {{?key?}} (dynamic list) in an existing template.

    path:   source .hwpx or .docx template
    keys:   placeholder name(s) to convert (single string or list)
    output: destination path; if None, overwrites the source file in-place

    Returns the output path.
    """
    path = Path(path)
    output_path = Path(output) if output is not None else path

    if isinstance(keys, str):
        keys = [keys]
    if not keys:
        return output_path

    ext = path.suffix.lower()
    if ext == ".hwpx":
        _convert_hwpx_placeholders(path, keys, output_path)
    elif ext == ".docx":
        _convert_docx_placeholders(path, keys, output_path)
    else:
        raise DocPilotError(
            f"Unsupported format '{ext}'",
            detail="Supported: .hwpx, .docx",
        )
    return output_path


def _convert_hwpx_placeholders(path: Path, keys: list[str], output: Path) -> None:
    import io

    with zipfile.ZipFile(path, "r") as zf_in:
        names = zf_in.namelist()
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf_out:
            for name in names:
                data = zf_in.read(name)
                if name.endswith("section0.xml"):
                    text = data.decode("utf-8")
                    for key in keys:
                        text = re.sub(
                            r"\{\{" + re.escape(key) + r"\}\}",
                            "{{?" + key + "?}}",
                            text,
                        )
                    data = text.encode("utf-8")
                compress = zipfile.ZIP_STORED if name == "mimetype" else zipfile.ZIP_DEFLATED
                zf_out.writestr(name, data, compress_type=compress)
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_bytes(buf.getvalue())


def _convert_docx_placeholders(path: Path, keys: list[str], output: Path) -> None:
    try:
        import docx
    except ImportError as e:
        raise DocPilotError("python-docx required: pip install python-docx") from e

    pattern = re.compile(
        r"\{\{(" + "|".join(re.escape(k) for k in keys) + r")\}\}"
    )

    doc = docx.Document(str(path))

    def _convert_para(para) -> None:
        full_text = "".join(run.text for run in para.runs)
        if not pattern.search(full_text):
            return
        new_text = pattern.sub(lambda m: "{{?" + m.group(1) + "?}}", full_text)
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""

    for para in doc.paragraphs:
        _convert_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _convert_para(para)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def _infer_sections_from_content(content: str, mapper) -> list[str]:
    """Use LLM to infer fillable section names from a document's text content."""
    import json
    from docpilot.mapping.base import TemplateSection

    mapping = mapper.map(
        content=content[:4000],
        sections=[
            TemplateSection(
                name="sections",
                description=(
                    "위 문서 구조를 분석해 채워야 할 섹션 이름 목록을 추출하세요. "
                    "헤딩·제목·표 컬럼·항목명 등 구조적 요소를 기반으로 하되 "
                    "내용이 들어가야 할 곳의 레이블을 우선하세요. "
                    "결과는 JSON 배열 문자열로만 반환하세요: [\"섹션1\", \"섹션2\", ...]"
                ),
            )
        ],
    )
    raw = mapping.sections.get("sections", "[]")
    try:
        sections = json.loads(raw)
        if isinstance(sections, list):
            return [str(s) for s in sections if s]
    except (json.JSONDecodeError, ValueError):
        pass
    return []


def _apply_reference_mode(template_path: Path, mapper) -> tuple[Path, list]:
    """
    Infer section structure via LLM from a placeholder-less reference document, then
    inject {{placeholder}}s into a temp copy of it. Costs one extra LLM call (structure
    inference) beyond the normal content-mapping call.

    Returns (new_template_path, sections). Raises DocPilotError if the content can't be
    read or the LLM fails to infer any section names.
    """
    from docpilot.mapping.base import TemplateSection
    from docpilot.template_generator.generator import _build_template, _build_docx_template

    ref_content = _ingest_instructions_doc(template_path)
    if not ref_content:
        raise DocPilotError(
            "No {{placeholders}} found and could not read template content",
            detail=str(template_path),
        )
    placeholder_names = _infer_sections_from_content(ref_content, mapper)
    if not placeholder_names:
        raise DocPilotError(
            "No {{placeholders}} found and LLM could not infer sections",
            detail=str(template_path),
        )

    ref_ext = template_path.suffix.lower()
    tpl_suffix = ref_ext if ref_ext in (".hwpx", ".docx") else ".hwpx"
    tmp = tempfile.NamedTemporaryFile(suffix=tpl_suffix, delete=False)
    tmp.close()
    tmp_tpl = Path(tmp.name)
    atexit.register(lambda p=tmp_tpl: p.unlink(missing_ok=True))

    if ref_ext == ".docx":
        _build_docx_template(template_path, placeholder_names, tmp_tpl)
    elif ref_ext == ".hwpx":
        _build_template(template_path, placeholder_names, tmp_tpl)
    else:
        # Non-buildable format (PDF, TXT …): use built-in report as base
        _build_template(_assemble_builtin_hwpx("report"), placeholder_names, tmp_tpl)

    return tmp_tpl, [TemplateSection(name=s) for s in placeholder_names]


def _sidecar_source(template: str | Path) -> Path:
    """
    Path to resolve sidecar.json against.

    Built-in templates are resolved to a temp-assembled HWPX copy for actual reading/
    building (see _assemble_builtin_hwpx), but their sidecar.json lives in the original
    templates/<name>/ directory — so sidecar lookup must use *template*, not the resolved
    template_path, or it silently finds nothing.
    """
    path = Path(template)
    if path.exists():
        return path
    if (_BUILTIN_TEMPLATES / str(template)).is_dir():
        return _BUILTIN_TEMPLATES / str(template)
    return path


def _resolve_template_path(template: str | Path, embed_fn=None) -> tuple[Path, dict]:
    """Return (template_path, sections_meta). sections_meta is non-empty only for DB templates."""
    path = Path(template)
    if path.exists():
        return path, {}

    name = str(template)
    if (_BUILTIN_TEMPLATES / name).is_dir():
        return _assemble_builtin_hwpx(name), {}

    for ext in (".hwpx", ".docx", ".pdf"):
        candidate = _BUILTIN_TEMPLATES / f"{template}{ext}"
        if candidate.exists():
            return candidate, {}
        local = Path("templates") / f"{template}{ext}"
        if local.exists():
            return local, {}

    # DB lookup by name
    try:
        from docpilot.db import template_store
        records = template_store.search(name, embed_fn=embed_fn, top_k=1, fallback=False)
        if records:
            record = records[0]
            section_xml_path = Path(record.path)
            if section_xml_path.exists():
                from docpilot.builder.hwpx_dynamic_builder import pack_hwpx
                _base_header = _BUILTIN_TEMPLATES / "base" / "Contents" / "header.xml"
                header = (
                    Path(record.header_xml)
                    if record.header_xml and Path(record.header_xml).exists()
                    else _base_header
                )
                tmp = tempfile.NamedTemporaryFile(suffix=".hwpx", delete=False)
                tmp.close()
                tmp_hwpx = Path(tmp.name)
                atexit.register(lambda p=tmp_hwpx: p.unlink(missing_ok=True))
                pack_hwpx(header, section_xml_path.read_text(encoding="utf-8"), tmp_hwpx)
                sections_meta = (record.metadata_ or {}).get("sections", {})
                return tmp_hwpx, sections_meta
    except Exception:
        pass

    raise DocPilotError("Template not found", detail=str(template))


def _resolve_builder(output: Path):
    from docpilot.builder import HwpxBuilder, PdfBuilder, DocxBuilder

    match output.suffix.lower():
        case ".hwpx":
            return HwpxBuilder()
        case ".pdf":
            return PdfBuilder()
        case ".docx":
            return DocxBuilder()
        case _:
            raise BuilderError(
                f"Unsupported output format '{output.suffix}'",
                detail="Supported: .hwpx, .pdf, .docx",
            )


def _merge_section_meta(sections: list, template_path: Path, sidecar, db_sections_meta: dict) -> list:
    """Apply sidecar.json rules, DB sections_meta, and template style hints onto *sections* in place."""
    if sidecar is not None and sidecar.sections:
        _sc_map = {s.name: s for s in sidecar.sections}
        for _sec in sections:
            _sc = _sc_map.get(_sec.name)
            if _sc is None:
                continue
            if not _sec.description and _sc.description:
                _sec.description = _sc.description
            if not _sec.rule and _sc.rule:
                _sec.rule = _sc.rule
            if not _sec.style_hint and _sc.style_hint:
                _sec.style_hint = _sc.style_hint
            if _sc.is_list:
                _sec.is_list = True
            if _sc.optional:
                _sec.optional = True
            if _sc.group_max > 0 and _sec.group_max == 0:
                _sec.group_max = _sc.group_max

    if db_sections_meta:
        for _sec in sections:
            _dm = db_sections_meta.get(_sec.name, {})
            if not _dm:
                continue
            if not _sec.description and _dm.get("description"):
                _sec.description = _dm["description"]
            if not _sec.rule and _dm.get("rule"):
                _sec.rule = _dm["rule"]
            if not _sec.style_hint and _dm.get("style_hint"):
                _sec.style_hint = _dm["style_hint"]
            if not _sec.is_list and _dm.get("is_list"):
                _sec.is_list = True
            if not _sec.optional and _dm.get("optional"):
                _sec.optional = True
            if _sec.group_max == 0 and _dm.get("group_max", 0) > 0:
                _sec.group_max = _dm["group_max"]

    style_hints: dict[str, str] = {}
    _tpl_ext = template_path.suffix.lower()
    if _tpl_ext == ".hwpx":
        from docpilot.builder.hwpx_analyzer import extract_style_hints
        style_hints = extract_style_hints(template_path)
    elif _tpl_ext == ".docx":
        from docpilot.builder.docx_analyzer import extract_style_hints
        style_hints = extract_style_hints(template_path)

    for s in sections:
        if s.style_hint:
            continue
        if s.is_list:
            for candidate in [f"{s.name}1", s.name]:
                hint = style_hints.get(candidate, "")
                if hint:
                    s.style_hint = hint
                    break
        else:
            s.style_hint = style_hints.get(s.name, "")

    return sections


def _resolve_sections(
    template: str | Path,
    template_path: Path,
    db_sections_meta: dict,
    mapper=None,
) -> tuple[Path, list, str]:
    """
    Extract + merge section metadata for a template.

    template:      the original name/path passed by the caller — used to locate sidecar.json,
                    since built-in templates resolve to a temp-assembled file (see _sidecar_source).
    template_path: the resolved, readable template file (from _resolve_template_path).
    mapper:        if given, a placeholder-less reference document triggers LLM-based section
                    inference (Reference Mode — one extra LLM call) instead of raising. If None
                    (default; used by describe_template()/fill_template(), which stay LLM-free),
                    raises DocPilotError instead.

    Returns (template_path, sections, instructions). template_path is reassigned only when
    Reference Mode ran (it then points at the generated temp template).
    """
    sections = _extract_placeholder_sections(template_path)
    if not sections:
        if mapper is None:
            raise DocPilotError(
                "템플릿에서 {{placeholder}}를 찾을 수 없습니다.",
                detail=(
                    f"{template_path} — 플레이스홀더가 없는 참조 문서는 generate_template()으로 "
                    "먼저 템플릿을 만들거나, LLM 추론을 지원하는 generate()/generate_from_content()를 "
                    "쓰세요."
                ),
            )
        template_path, sections = _apply_reference_mode(template_path, mapper)

    from docpilot.mapping.sidecar import load_sidecar
    sidecar = load_sidecar(_sidecar_source(template))
    sections = _merge_section_meta(sections, template_path, sidecar, db_sections_meta)
    instructions = sidecar.instructions if sidecar is not None else ""
    return template_path, sections, instructions


def describe_template(template: str | Path) -> dict:
    """
    Return a template's fillable structure — no LLM/RAG call.

    Use this before fill_template() to see exact section keys, per-section rules/
    descriptions, and a ready-to-copy example dict shape.
    """
    template_path, db_sections_meta = _resolve_template_path(template)
    template_path, sections, instructions = _resolve_sections(template, template_path, db_sections_meta)

    example: dict[str, str | list[str]] = {}
    for s in sections:
        if s.is_list:
            example[s.name] = ["항목1", "항목2"]
        elif s.optional:
            example[s.name] = ""
        else:
            example[s.name] = "..."

    return {
        "template": str(template_path),
        "instructions": instructions,
        "sections": [
            {
                "name": s.name,
                "description": s.description,
                "rule": s.rule,
                "style_hint": s.style_hint,
                "optional": s.optional,
                "is_list": s.is_list,
                "group_max": s.group_max,
            }
            for s in sections
        ],
        "example": example,
    }


def fill_template(
    template: str | Path,
    sections: dict[str, str | list[str]],
    output: str | Path,
) -> Path:
    """
    Fill an already-authored section dict into a template — pure mechanical build,
    no LLM/RAG call. Validates keys against the template's actual placeholders first,
    and warns if any placeholder is left unfilled in the output.
    """
    template_path, db_sections_meta = _resolve_template_path(template)
    template_path, resolved, _ = _resolve_sections(template, template_path, db_sections_meta)
    resolved_names = {s.name for s in resolved}

    required_missing = [
        s.name for s in resolved
        if not s.optional and not s.is_list and not str(sections.get(s.name, "")).strip()
    ]
    unknown_keys = [k for k in sections if k not in resolved_names]
    if required_missing or unknown_keys:
        detail_lines = []
        if required_missing:
            detail_lines.append(f"누락된 필수 섹션: {', '.join(required_missing)}")
        if unknown_keys:
            detail_lines.append(f"템플릿에 없는 키 (오탈자 의심): {', '.join(unknown_keys)}")
        raise MappingError(
            "fill_template 입력 검증 실패 — describe_template()으로 정확한 섹션 키를 확인하세요.",
            detail="\n".join(detail_lines),
        )

    fill: dict[str, str | list[str]] = {}
    for s in resolved:
        val = sections.get(s.name, [] if s.is_list else "")
        if s.is_list and isinstance(val, str):
            val = [v.strip() for v in val.split("\n") if v.strip()]
        fill[s.name] = val

    output_path = Path(output)
    builder = _resolve_builder(output_path)
    out_path = builder.build(template_path, fill, output_path)

    if out_path.suffix.lower() == ".hwpx":
        _validate_hwpx(out_path)

    leftover = _extract_placeholders(out_path)
    if leftover:
        import warnings
        warnings.warn(
            f"생성된 문서에 채워지지 않은 플레이스홀더가 남아 있습니다: {', '.join(leftover)}",
            stacklevel=2,
        )

    return out_path


class DocPilot:
    """
    Main entry point for docpilot.

    pilot = DocPilot(llm="claude")
    pilot.index(data_folder="./data")
    pilot.generate(data_folder="./data", template="research_report", output="./out.hwpx")
    """

    def __init__(
        self,
        llm: str | None = None,
        api_key: str | None = None,
        model: str | None = None,
        base_url: str | None = None,
        database_url: str | None = None,
        embed_fn=None,
        use_reranker: bool = False,
    ) -> None:
        self._llm = llm or os.environ.get("DOCPILOT_LLM", "claude")
        self._api_key = api_key
        if embed_fn is None:
            from docpilot.search.embedding import default_embed_fn
            embed_fn = default_embed_fn()  # None when sentence-transformers not installed
        self._embed_fn = embed_fn
        base_mapper = self._build_mapper(self._llm, api_key, model, base_url)

        from docpilot.mapping.rag import RagMapper
        self._mapper = base_mapper
        self._rag_mapper = RagMapper(base_mapper, embed_fn=embed_fn, use_reranker=use_reranker)

        from docpilot.db import client as db_client
        db_client.init(database_url)
        db_client.create_tables()

    def index(self, data_folder: str | Path, collection: str | None = None) -> list[int]:
        """Ingest and index all supported files in a folder.

        collection: optional tag to scope later searches (see SearchFilter.collection).
        """
        from docpilot.db import indexer
        return indexer.index_folder(data_folder, embed_fn=self._embed_fn, collection=collection)

    def search(
        self,
        query: str,
        top_k: int = 10,
        mode: str = "hybrid",
        filters=None,
        or_fallback: bool = True,
        highlight: bool = False,
        group_by_doc: bool = False,
        collection: str | None = None,
    ):
        """
        Search indexed documents.

        mode:
            ``"hybrid"``  — BM25 + Vector fused with RRF (default; falls back to BM25 when no embed_fn).
            ``"bm25"``    — Morpheme FTS5 / BM25 only.
            ``"vector"``  — Vector similarity only (embed_fn required).
            ``"exact"``   — ILIKE keyword match.
        filters:
            ``SearchFilter`` instance for source/MIME/metadata/date constraints.
        or_fallback:
            For BM25/hybrid — retry with OR logic when AND returns no results.
        highlight:
            If True, populate each result's ``.highlights`` spans for *query* terms.
            Render the marked-up text with ``docpilot.search.highlight.render(result)``.
        group_by_doc:
            If True, aggregate chunk-level results into per-document ``DocumentResult``
            objects — returns ``list[DocumentResult]`` instead of ``list[SearchResult]``.
        collection:
            Convenience shorthand for ``filters.collection`` — scopes the search to documents
            indexed with this tag. Raises ``ValueError`` if *filters* already sets a different
            ``collection``.

        "bm25"/"hybrid" modes fall back to "exact" and emit a ``UserWarning`` when
        kiwipiepy isn't installed, instead of raising.
        """
        import warnings
        from dataclasses import replace as _dc_replace
        from docpilot.exceptions import SearchError
        from docpilot.search import exact, morpheme
        from docpilot.search.models import SearchFilter

        if collection is not None:
            if filters is None:
                filters = SearchFilter(collection=collection)
            elif filters.collection is None:
                filters = _dc_replace(filters, collection=collection)
            elif filters.collection != collection:
                raise ValueError(
                    f"collection={collection!r} conflicts with filters.collection={filters.collection!r}"
                )

        try:
            match mode:
                case "exact":
                    results = exact.search(query, top_k=top_k, filters=filters)
                case "bm25":
                    results = morpheme.search(query, top_k=top_k, or_fallback=or_fallback, filters=filters)
                case "vector":
                    if self._embed_fn is None:
                        raise SearchError(
                            "vector 모드는 embed_fn 설정이 필요합니다",
                            detail="DOCPILOT_EMBED 환경변수 또는 DocPilot(embed_fn=...) 설정을 확인하세요",
                        )
                    from docpilot.search import embedding
                    results = embedding.search(query, embed_fn=self._embed_fn, top_k=top_k, filters=filters)
                case _:  # "hybrid"
                    from docpilot.search.hybrid import hybrid
                    results = hybrid(
                        query,
                        embed_fn=self._embed_fn,
                        top_k=top_k,
                        filters=filters,
                        or_fallback=or_fallback,
                    )
        except SearchError as e:
            if mode in ("bm25", "hybrid") and "kiwipiepy" in str(e):
                warnings.warn(
                    f"kiwipiepy 미설치로 exact 검색으로 대체합니다 (형태소 검색을 쓰려면 "
                    f"pip install kiwipiepy — 기본 설치에 포함되므로 보통은 발생하지 않는 상황입니다): {e}",
                    stacklevel=2,
                )
                results = exact.search(query, top_k=top_k, filters=filters)
            else:
                raise

        if highlight:
            from docpilot.search import highlight as highlight_fn
            results = [highlight_fn(r, query) for r in results]

        if group_by_doc:
            from docpilot.search import group_by_document
            return group_by_document(results, top_chunks=3, score="max")

        return results

    def generate(
        self,
        data_folder: str | Path,
        template: str | Path,
        output: str | Path,
        reindex: bool = False,
        extra_instructions: str | None = None,
        instructions_doc: str | Path | None = None,
        top_k: int = 10,
        collection: str | None = None,
    ) -> GenerateResult:
        """
        Full pipeline: index → search → map → build.

        template:           file path (.hwpx/.docx/.pdf) or built-in template name.
        output:             destination file path — extension determines builder.
        reindex:            re-index data_folder even if already indexed.
        extra_instructions: additional writing guidelines injected into the LLM prompt.
                            Use this to pass document-specific rules such as proposal
                            writing guidelines extracted from an RFP.
        collection:         optional tag applied when indexing data_folder. When given, RAG
                            retrieval is scoped by this tag instead of the data_folder path,
                            so it keeps working even if files are later moved.
        """
        template_path, _db_sections_meta = self._resolve_template(template)
        output_path = Path(output)

        if template_path.suffix.lower() == ".docx" and output_path.suffix.lower() == ".hwpx":
            from docpilot.builder.hwp_convert import convert_to_hwpx
            _tmp = tempfile.NamedTemporaryFile(suffix=".hwpx", delete=False)
            _tmp.close()
            _tmp_path = Path(_tmp.name)
            atexit.register(lambda p=_tmp_path: p.unlink(missing_ok=True))
            template_path = convert_to_hwpx(template_path, _tmp_path)

        from docpilot.mapping.sidecar import load_sidecar
        _sidecar = load_sidecar(_sidecar_source(template))
        extra_instructions = _resolve_instructions(
            extra_instructions,
            _sidecar.instructions if _sidecar is not None else "",
            instructions_doc,
        )

        from docpilot.db import indexer
        doc_ids = indexer.index_folder(
            data_folder, embed_fn=self._embed_fn, force=reindex, collection=collection
        )
        if not doc_ids:
            raise DocPilotError(
                "데이터 폴더에서 인덱싱된 문서가 없습니다.",
                detail=(
                    f"{data_folder} 에서 지원하는 파일을 찾지 못했거나 "
                    "모든 파일의 수집(ingestion)에 실패했습니다. "
                    "지원 형식: .txt .md .csv .hwpx .hwp .docx .pdf"
                ),
            )

        template_path, sections, _ = _resolve_sections(
            template, template_path, _db_sections_meta, mapper=self._mapper,
        )

        from docpilot.search.models import SearchFilter
        if collection is not None:
            filters = SearchFilter(collection=collection)
        else:
            filters = SearchFilter(source_pattern=f"{Path(data_folder).resolve()}{os.sep}*")
        mapping_result = self._rag_mapper.map(sections, instructions=extra_instructions, top_k=top_k, filters=filters)

        builder = self._build_builder(output_path)
        out_path = builder.build(template_path, mapping_result.sections, output_path)

        if out_path.suffix.lower() == ".hwpx":
            _validate_hwpx(out_path)

        return GenerateResult(
            path=out_path,
            model=mapping_result.model,
            input_tokens=mapping_result.input_tokens,
            output_tokens=mapping_result.output_tokens,
            elapsed_seconds=mapping_result.elapsed_seconds,
        )

    def generate_from_content(
        self,
        content: str | list[str | Path] | list[IngestedDocument],
        template: str | Path,
        output: str | Path,
        extra_instructions: str | None = None,
        instructions_doc: str | Path | None = None,
        max_input_tokens: int | None = None,
    ) -> GenerateResult:
        """
        Fill a template from already-prepared content — no RAG retrieval, no data folder.

        content:            a finished string, a list of IngestedDocuments, or a list of file
                             paths (ingested directly, no chunking/embedding/DB storage).
        template:            {{placeholder}} template, or a placeholder-less reference document —
                             the latter costs one extra LLM call to infer section structure
                             (Reference Mode, same as generate()).
        max_input_tokens:    if set, raises MappingError when the estimated input token count
                             exceeds it. If unset, only warns — RAG's top_k has no equivalent
                             here, so there's no automatic context-size limit.
        """
        template_path, db_sections_meta = self._resolve_template(template)
        output_path = Path(output)

        if template_path.suffix.lower() == ".docx" and output_path.suffix.lower() == ".hwpx":
            from docpilot.builder.hwp_convert import convert_to_hwpx
            _tmp = tempfile.NamedTemporaryFile(suffix=".hwpx", delete=False)
            _tmp.close()
            _tmp_path = Path(_tmp.name)
            atexit.register(lambda p=_tmp_path: p.unlink(missing_ok=True))
            template_path = convert_to_hwpx(template_path, _tmp_path)

        template_path, sections, sidecar_instructions = _resolve_sections(
            template, template_path, db_sections_meta, mapper=self._mapper,
        )
        extra_instructions = _resolve_instructions(extra_instructions, sidecar_instructions, instructions_doc)

        content_str = _resolve_content(content)
        _check_content_size(content_str, max_input_tokens)

        mapping_result = self._mapper.map(content_str, sections, instructions=extra_instructions)

        builder = self._build_builder(output_path)
        out_path = builder.build(template_path, mapping_result.sections, output_path)

        if out_path.suffix.lower() == ".hwpx":
            _validate_hwpx(out_path)

        return GenerateResult(
            path=out_path,
            model=mapping_result.model,
            input_tokens=mapping_result.input_tokens,
            output_tokens=mapping_result.output_tokens,
            elapsed_seconds=mapping_result.elapsed_seconds,
        )

    def generate_template(
        self,
        samples: list[str | Path],
        output: str | Path,
        use_llm: bool | None = None,
    ) -> Path:
        """Generate an HWPX template from sample documents."""
        from docpilot.template_generator import generate
        return generate(
            samples=samples,
            output=output,
            use_llm=use_llm,
            llm_mapper=self._mapper if use_llm else None,
        )

    def save_template(
        self,
        name: str,
        path: str | Path,
        description: str = "",
        tags: list[str] | None = None,
        sections_meta: dict | None = None,
        auto_sections_meta: bool = True,
    ) -> int:
        """
        Save an HWPX template to the DB so it can be found by name in generate().

        Extracts section0.xml and header.xml from the .hwpx ZIP into
        ~/.docpilot/templates/<name>/ for persistent access across sessions.

        auto_sections_meta: if True (default) and sections_meta is None, ask the LLM to
            infer per-section description/rule for each placeholder — this makes a
            billed API call every time save_template() is called without explicit
            sections_meta. Pass auto_sections_meta=False or sections_meta={} to skip it
            (a UserWarning is also raised when the automatic call is about to run).

        Returns the saved template record ID.
        """
        import zipfile
        from docpilot.db import template_store

        path = Path(path)
        if not path.exists():
            raise DocPilotError("Template file not found", detail=str(path))
        if path.suffix.lower() != ".hwpx":
            raise DocPilotError(
                "save_template() only supports .hwpx files",
                detail=f"Got '{path.suffix}'. Pass .docx/.pdf paths directly to generate().",
            )

        store_dir = Path.home() / ".docpilot" / "templates" / name
        store_dir.mkdir(parents=True, exist_ok=True)
        section_xml_path = store_dir / "section0.xml"
        header_xml_path = store_dir / "header.xml"

        with zipfile.ZipFile(path, "r") as zf:
            names = zf.namelist()
            section_candidates = [n for n in names if n.endswith("section0.xml")]
            if not section_candidates:
                raise DocPilotError("No section0.xml in HWPX", detail=str(path))
            section_xml_path.write_bytes(zf.read(section_candidates[0]))
            if "Contents/header.xml" in names:
                header_xml_path.write_bytes(zf.read("Contents/header.xml"))
            else:
                header_xml_path = None

        return template_store.save(
            name=name,
            path=str(section_xml_path),
            description=description or name,
            header_xml=str(header_xml_path) if header_xml_path else None,
            tags=tags,
            sections_meta=sections_meta,
            auto_sections_meta=auto_sections_meta,
            mapper=self._mapper,
            embed_fn=self._embed_fn,
        )

    def estimate_cost(
        self,
        data_folder: str | Path,
        template: str | Path,
        quick: bool = False,
        collection: str | None = None,
    ) -> str:
        """
        Estimate API token cost before generating.

        quick=True (default in MCP): file-size based estimate, no indexing or API call.
        quick=False: full indexing + RAG retrieval + token-counting API for accuracy.
        """
        template_path, db_sections_meta = self._resolve_template(template)
        template_path, sections, _ = _resolve_sections(
            template, template_path, db_sections_meta, mapper=self._mapper,
        )

        model: str = getattr(self._mapper, "_model", "claude-sonnet-4-6")
        in_price, out_price = _MODEL_PRICING.get(model, (3.00, 15.00))
        est_output = len(sections) * _EST_OUTPUT_TOKENS_PER_SECTION

        if quick:
            input_tokens = _estimate_tokens_from_folder(Path(data_folder), len(sections))
            input_cost = input_tokens / 1_000_000 * in_price
            output_cost = est_output / 1_000_000 * out_price
            lines = [
                "=== docpilot 비용 추정 (빠른 추정) ===",
                f"모델:             {model}",
                f"섹션 수:          {len(sections)}개",
                f"입력 토큰 (추정): {input_tokens:,}  (파일 크기 기반, ±30% 오차)",
                f"출력 토큰 (추정): {est_output:,}  (섹션당 {_EST_OUTPUT_TOKENS_PER_SECTION} 추정)",
                f"예상 비용:        ${input_cost + output_cost:.4f}",
                f"  입력 ${in_price:.2f}/1M  →  ${input_cost:.4f}",
                f"  출력 ${out_price:.2f}/1M  →  ${output_cost:.4f}",
                "",
                "정확한 추정: estimate_cost(quick=False) — 인덱싱 + 토큰 카운팅 API 사용",
            ]
            return "\n".join(lines)

        from docpilot.db import indexer
        indexer.index_folder(data_folder, embed_fn=self._embed_fn, collection=collection)
        content = self._rag_mapper.retrieve_content(sections)

        if not hasattr(self._mapper, "count_tokens"):
            n = len(sections)
            return (
                f"섹션 수: {n}개\n"
                f"토큰 카운팅은 Claude 매퍼에서만 지원됩니다.\n"
                f"섹션당 ~3,000 입력 + ~{_EST_OUTPUT_TOKENS_PER_SECTION} 출력 토큰 기준\n"
                f"대략 {n * 3000:,} 입력 / {n * _EST_OUTPUT_TOKENS_PER_SECTION:,} 출력 예상"
            )

        input_tokens = self._mapper.count_tokens(content, sections)
        input_cost = input_tokens / 1_000_000 * in_price
        output_cost = est_output / 1_000_000 * out_price

        lines = [
            "=== docpilot 비용 추정 ===",
            f"모델:             {model}",
            f"섹션 수:          {len(sections)}개",
            f"입력 토큰:        {input_tokens:,}",
            f"출력 토큰 (추정): {est_output:,}  (섹션당 {_EST_OUTPUT_TOKENS_PER_SECTION} 추정)",
            f"예상 비용:        ${input_cost + output_cost:.4f}",
            f"  입력 ${in_price:.2f}/1M  →  ${input_cost:.4f}",
            f"  출력 ${out_price:.2f}/1M  →  ${output_cost:.4f}",
        ]
        return "\n".join(lines)

    def estimate_cost_from_content(
        self,
        content: str | list[str | Path] | list[IngestedDocument],
        template: str | Path,
        quick: bool = False,
    ) -> str:
        """
        Estimate API token cost for generate_from_content() — no indexing, no RAG retrieval.

        quick=True: byte-based heuristic, no API call.
        quick=False (default): real token-counting API call against the actual content.
        """
        template_path, db_sections_meta = self._resolve_template(template)
        template_path, sections, _ = _resolve_sections(
            template, template_path, db_sections_meta, mapper=self._mapper,
        )
        content_str = _resolve_content(content)

        model: str = getattr(self._mapper, "_model", "claude-sonnet-4-6")
        in_price, out_price = _MODEL_PRICING.get(model, (3.00, 15.00))
        est_output = len(sections) * _EST_OUTPUT_TOKENS_PER_SECTION

        if quick:
            input_tokens = _estimate_tokens_raw(content_str)
            input_cost = input_tokens / 1_000_000 * in_price
            output_cost = est_output / 1_000_000 * out_price
            lines = [
                "=== docpilot 비용 추정 (빠른 추정, RAG 없음) ===",
                f"모델:             {model}",
                f"섹션 수:          {len(sections)}개",
                f"입력 토큰 (추정): {input_tokens:,}  (바이트 기반, ±30% 오차)",
                f"출력 토큰 (추정): {est_output:,}  (섹션당 {_EST_OUTPUT_TOKENS_PER_SECTION} 추정)",
                f"예상 비용:        ${input_cost + output_cost:.4f}",
                f"  입력 ${in_price:.2f}/1M  →  ${input_cost:.4f}",
                f"  출력 ${out_price:.2f}/1M  →  ${output_cost:.4f}",
                "",
                "정확한 추정: estimate_cost_from_content(quick=False) — 토큰 카운팅 API 사용",
            ]
            return "\n".join(lines)

        if not hasattr(self._mapper, "count_tokens"):
            n = len(sections)
            return (
                f"섹션 수: {n}개\n"
                f"토큰 카운팅은 Claude 매퍼에서만 지원됩니다.\n"
                f"섹션당 ~3,000 입력 + ~{_EST_OUTPUT_TOKENS_PER_SECTION} 출력 토큰 기준\n"
                f"대략 {n * 3000:,} 입력 / {n * _EST_OUTPUT_TOKENS_PER_SECTION:,} 출력 예상"
            )

        input_tokens = self._mapper.count_tokens(content_str, sections)
        input_cost = input_tokens / 1_000_000 * in_price
        output_cost = est_output / 1_000_000 * out_price

        lines = [
            "=== docpilot 비용 추정 (RAG 없음) ===",
            f"모델:             {model}",
            f"섹션 수:          {len(sections)}개",
            f"입력 토큰:        {input_tokens:,}",
            f"출력 토큰 (추정): {est_output:,}  (섹션당 {_EST_OUTPUT_TOKENS_PER_SECTION} 추정)",
            f"예상 비용:        ${input_cost + output_cost:.4f}",
            f"  입력 ${in_price:.2f}/1M  →  ${input_cost:.4f}",
            f"  출력 ${out_price:.2f}/1M  →  ${output_cost:.4f}",
        ]
        return "\n".join(lines)

    def benchmark(
        self,
        data_folder: str | Path,
        template: str | Path,
        output: str | Path,
        mappers: dict | None = None,
        collection: str | None = None,
    ) -> str:
        """Run mapping benchmark across multiple LLM mappers and return report."""
        from docpilot.mapping import benchmark, ClaudeMapper, OpenAIMapper

        template_path, db_sections_meta = self._resolve_template(template)
        self.index(data_folder, collection=collection)

        template_path, template_sections, _ = _resolve_sections(
            template, template_path, db_sections_meta, mapper=self._mapper,
        )
        content = self._rag_mapper.retrieve_content(template_sections)

        if mappers is None:
            mappers = {"claude": ClaudeMapper(api_key=self._api_key)}
            oai_key = os.environ.get("OPENAI_API_KEY")
            if oai_key:
                mappers["openai"] = OpenAIMapper(api_key=oai_key)

        results = benchmark.run(content, template_sections, mappers)
        return benchmark.report(results)

    def benchmark_from_content(
        self,
        content: str | list[str | Path] | list[IngestedDocument],
        template: str | Path,
        output: str | Path,
        mappers: dict | None = None,
        max_input_tokens: int | None = None,
    ) -> str:
        """Run mapping benchmark across multiple LLM mappers — no indexing, no RAG retrieval."""
        from docpilot.mapping import benchmark, ClaudeMapper, OpenAIMapper

        template_path, db_sections_meta = self._resolve_template(template)
        template_path, template_sections, _ = _resolve_sections(
            template, template_path, db_sections_meta, mapper=self._mapper,
        )
        content_str = _resolve_content(content)
        _check_content_size(content_str, max_input_tokens)

        if mappers is None:
            mappers = {"claude": ClaudeMapper(api_key=self._api_key)}
            oai_key = os.environ.get("OPENAI_API_KEY")
            if oai_key:
                mappers["openai"] = OpenAIMapper(api_key=oai_key)

        results = benchmark.run(content_str, template_sections, mappers)
        return benchmark.report(results)

    def _resolve_template(self, template: str | Path) -> tuple[Path, dict]:
        """Return (template_path, sections_meta). sections_meta is non-empty only for DB templates."""
        return _resolve_template_path(template, self._embed_fn)

    @staticmethod
    def list_templates() -> dict[str, str]:
        """Return available template names and their descriptions (built-in + DB-saved)."""
        result = _get_builtin_metadata()
        try:
            from docpilot.db import template_store
            for record in template_store.list_all():
                if record.name not in result:
                    result[record.name] = record.description
        except Exception:
            pass
        return result

    @staticmethod
    def delete_template(template_id: int) -> bool:
        """Delete a saved template by ID. Returns True if found and deleted."""
        from docpilot.db import template_store
        return template_store.delete(template_id)

    @staticmethod
    def delete_template_by_path(path: str | Path) -> bool:
        """Delete a saved template by its section0.xml path. Returns True if found and deleted."""
        from docpilot.db import template_store
        return template_store.delete_by_path(path)

    @staticmethod
    def convert_to_list_placeholder(
        path: str | Path,
        keys: list[str] | str,
        output: str | Path | None = None,
    ) -> Path:
        """
        Convert {{key}} placeholders to {{?key?}} (dynamic list) in a template file.

        keys:   placeholder name(s) to convert
        output: destination path; if None, overwrites source in-place
        """
        return convert_to_list_placeholder(path, keys, output)

    @staticmethod
    def suggest_extras(folder: str | Path) -> dict:
        """Scan a data folder and suggest which docpilot extras to install."""
        return suggest_extras(folder)

    @staticmethod
    def _build_mapper(llm: str, api_key: str | None, model: str | None, base_url: str | None):
        from docpilot.mapping import ClaudeMapper, OpenAIMapper, GeminiMapper
        from docpilot.mapping.openai_compat import GrokMapper, OllamaMapper

        match llm.lower():
            case "claude":
                return ClaudeMapper(api_key=api_key, **(_m(model)))
            case "openai":
                return OpenAIMapper(api_key=api_key, **(_m(model)))
            case "gemini":
                return GeminiMapper(api_key=api_key, **(_m(model)))
            case "grok":
                return GrokMapper(api_key=api_key, **(_m(model)))
            case "ollama":
                return OllamaMapper(**(_m(model)), **(_b(base_url)))
            case _:
                raise MappingError(
                    f"Unknown LLM '{llm}'. "
                    "Supported: claude, openai, gemini, grok, ollama"
                )

    @staticmethod
    def _build_builder(output: Path):
        return _resolve_builder(output)
